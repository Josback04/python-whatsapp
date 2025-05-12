# app/services/docx_service.py

import os
import re
import logging # Import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# Import the function that uploads to Drive
from app.services.drive_service import upload_file_to_drive 

def generate_docx(user_name, ai_response):
    """
    Génère un fichier Word (DOCX), l'upload sur Google Drive, 
    et retourne l'URL Google Drive ou None en cas d'échec.
    """
    # Utiliser un nom de fichier temporaire ou spécifique
    # Il est préférable de le créer dans un dossier temporaire si possible, 
    # mais pour la simplicité, on le met dans le répertoire courant.
    file_name = f"{user_name}_plan_daffaire_{os.urandom(4).hex()}.docx" 
    file_path = os.path.join(file_name) # Chemin complet du fichier local temporaire
    
    logging.info(f"Génération du document local : {file_path}")
    
    drive_url = None # Initialiser l'URL retournée à None

    try:
        doc = Document()
        
        # Ajouter un titre
        # S'assurer que le style 'Heading 1' existe ou utiliser un style par défaut/le créer
        try:
             # Utiliser 'Heading 1' si disponible, sinon 'Title' ou un style de base
             title_style = 'Heading 1' if 'Heading 1' in doc.styles else 'Title'
             title = doc.add_paragraph(f"Recommandation pour {user_name}", style=title_style)
        except KeyError: # Au cas où même 'Title' n'existe pas
             title = doc.add_paragraph(f"Recommandation pour {user_name}")
             # Appliquer manuellement le formatage si nécessaire (gras, taille)
             title.runs[0].font.bold = True
             title.runs[0].font.size = Pt(16) # Exemple
             
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Traiter chaque paragraphe de la réponse
        for paragraph in ai_response.split("\n"):
            if paragraph.strip() == "":
                p = doc.add_paragraph() # Ajouter un paragraphe vide pour les sauts de ligne
                continue
                
            bold_parts = re.split(r"(\*\*.*?\*\*)", paragraph)
            
            p = doc.add_paragraph()
            for part in bold_parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part.strip("**"))
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Enregistrer le document localement
        doc.save(file_path)
        logging.info(f"Document local sauvegardé : {file_path}")

        # Appeler la fonction d'upload et capturer l'URL retournée
        logging.info(f"Tentative d'upload de {file_path} vers Google Drive...")
        drive_url = upload_file_to_drive(file_path) # <--- CAPTURER LE RETOUR ICI

        if drive_url:
             logging.info(f"Upload vers Google Drive réussi. URL: {drive_url}")
        else:
             logging.error(f"Échec de l'upload de {file_path} vers Google Drive.")
             # drive_url est déjà None

    except Exception as e:
        logging.exception(f"Erreur lors de la génération ou de l'upload du DOCX pour {user_name}: {e}")
        drive_url = None # S'assurer que l'URL est None en cas d'erreur ici aussi
        
    finally:
        # Nettoyer le fichier local, qu'il y ait eu succès ou échec d'upload
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                logging.info(f"Fichier local temporaire supprimé : {file_path}")
            except OSError as e:
                logging.error(f"Erreur lors de la suppression du fichier local {file_path}: {e}")

    # Retourner l'URL Google Drive (qui sera None si une erreur est survenue)
    return drive_url # <--- RETOURNER L'URL CAPTURÉE (ou None)